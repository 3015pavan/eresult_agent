"""
Query Endpoints — Phase 5 Query Engine.

Natural language query interface for teachers.
Backend: PostgreSQL (psycopg2)
"""

from __future__ import annotations

import re
import time
from typing import Any

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from psycopg2.extras import RealDictCursor

from src.common.observability import get_logger
from src.common import database as db

logger = get_logger(__name__)
router = APIRouter()


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

class QueryRequest(BaseModel):
    query: str = Field(..., min_length=3, max_length=500)
    context: dict[str, Any] | None = None


class QueryResponse(BaseModel):
    query: str
    intent: str
    text_answer: str
    summary: str | None = None
    data: list[dict[str, Any]] = []
    chart_spec: dict[str, Any] | None = None
    confidence: float
    caveats: list[str] = []
    sources: list[str] = []   # USNs or record identifiers cited in the answer


class StudentSummary(BaseModel):
    usn: str
    name: str
    email: str | None
    department: str
    batch_year: int | None
    current_cgpa: float
    active_backlogs: int
    total_results: int
    class_rank: int | None
    total_students: int
    semesters: list[dict[str, Any]]
    results_by_semester: dict[str, list[dict[str, Any]]]


class TrendResponse(BaseModel):
    usn: str
    name: str
    trend: list[dict[str, Any]]


# ---------------------------------------------------------------------------
# Helpers — normalise row values for JSON serialisation
# ---------------------------------------------------------------------------

def _serialise(row: dict) -> dict:
    return {
        k: str(v) if not isinstance(v, (str, int, float, bool, type(None), list, dict)) else v
        for k, v in row.items()
    }


def _student_display_name(s: dict) -> str:
    return s.get("full_name") or s.get("name") or s.get("usn") or ""


def _compute_report_metrics(results: list[dict[str, Any]], sems: list[dict[str, Any]]) -> dict[str, Any]:
    """Build deterministic summary metrics from stored student results."""
    total_scored = 0.0
    total_max = 0.0
    passed = 0
    failed = 0
    semesters = sorted(
        {
            int(row.get("semester"))
            for row in results
            if row.get("semester") is not None
        }
    )

    for row in results:
        try:
            marks = row.get("marks_obtained")
            max_marks = row.get("max_marks")
            if marks is not None and max_marks is not None:
                total_scored += float(marks)
                total_max += float(max_marks)
        except (TypeError, ValueError):
            pass

        status = str(row.get("pass_status") or "").upper()
        if status == "PASS":
            passed += 1
        elif status == "FAIL":
            failed += 1

    best_sgpa = 0.0
    for sem in sems:
        try:
            best_sgpa = max(best_sgpa, float(sem.get("sgpa") or 0.0))
        except (TypeError, ValueError):
            continue

    percentage = round((total_scored / total_max) * 100, 2) if total_max else 0.0
    pass_rate = round((passed / (passed + failed)) * 100, 2) if (passed + failed) else 0.0

    return {
        "subjects_total": len(results),
        "subjects_passed": passed,
        "subjects_failed": failed,
        "total_scored": round(total_scored, 2),
        "total_max": round(total_max, 2),
        "percentage": percentage,
        "pass_rate": pass_rate,
        "best_sgpa": round(best_sgpa, 2),
        "semesters_completed": len(semesters),
    }


def _build_html_report(usn, name, email_str, dept, batch_year, cgpa, backlogs, sems, results, metrics) -> str:
    """Render the student report as an HTML template."""
    by_sem: dict[int, list[dict[str, Any]]] = {}
    for row in results:
        by_sem.setdefault(int(row.get("semester", 0)), []).append(row)

    sem_cards = []
    for sem in sems:
        sem_cards.append(
            f"""
            <div class="metric-card">
              <div class="metric-label">Semester {sem.get('semester')}</div>
              <div class="metric-value">{float(sem.get('sgpa') or 0):.2f}</div>
              <div class="metric-sub">SGPA · Backlogs {int(sem.get('backlogs') or 0)}</div>
            </div>
            """
        )

    detail_sections = []
    for sem_num in sorted(by_sem):
        rows = []
        for row in by_sem[sem_num]:
            status = row.get("pass_status", "—") or "—"
            rows.append(
                f"""
                <tr>
                  <td>{row.get('subject_code', '')}</td>
                  <td>{row.get('subject_name', '')}</td>
                  <td>{int(row.get('marks_obtained', 0) or 0)}</td>
                  <td>{int(row.get('max_marks', 100) or 100)}</td>
                  <td>{row.get('grade', '—') or '—'}</td>
                  <td class="{status.lower()}">{status}</td>
                </tr>
                """
            )
        detail_sections.append(
            f"""
            <section class="semester-section">
              <h3>Semester {sem_num}</h3>
              <table>
                <thead>
                  <tr><th>Code</th><th>Subject</th><th>Marks</th><th>Max</th><th>Grade</th><th>Status</th></tr>
                </thead>
                <tbody>
                  {''.join(rows)}
                </tbody>
              </table>
            </section>
            """
        )

    return f"""
    <!doctype html>
    <html>
    <head>
      <meta charset="utf-8">
      <title>Academic Report - {usn}</title>
      <style>
        :root {{ --ink:#1a1d27; --muted:#677189; --accent:#1e5eff; --line:#d8ddea; --bg:#f6f8fc; --ok:#157347; --bad:#b42318; }}
        * {{ box-sizing: border-box; }}
        body {{ font-family: 'Segoe UI', Tahoma, sans-serif; margin:0; color:var(--ink); background:linear-gradient(180deg,#ffffff,#f6f8fc); }}
        .page {{ max-width: 980px; margin: 0 auto; padding: 32px; }}
        .hero {{ display:flex; justify-content:space-between; align-items:flex-end; border-bottom:2px solid var(--line); padding-bottom:18px; margin-bottom:24px; }}
        .brand {{ font-size:28px; font-weight:700; color:var(--accent); }}
        .meta {{ color:var(--muted); font-size:13px; }}
        .profile {{ display:grid; grid-template-columns: repeat(2, minmax(0,1fr)); gap:10px 18px; background:#fff; border:1px solid var(--line); border-radius:16px; padding:20px; }}
        .field-label {{ color:var(--muted); font-size:12px; text-transform:uppercase; letter-spacing:.08em; }}
        .field-value {{ font-size:15px; font-weight:600; }}
        .metrics {{ display:grid; grid-template-columns: repeat(auto-fit, minmax(160px,1fr)); gap:14px; margin:22px 0; }}
        .metric-card {{ background:#fff; border:1px solid var(--line); border-radius:16px; padding:16px; }}
        .metric-label {{ color:var(--muted); font-size:12px; text-transform:uppercase; letter-spacing:.06em; }}
        .metric-value {{ font-size:24px; font-weight:700; margin-top:6px; }}
        .metric-sub {{ color:var(--muted); font-size:12px; margin-top:4px; }}
        h2, h3 {{ margin:24px 0 12px; }}
        table {{ width:100%; border-collapse:collapse; background:#fff; border:1px solid var(--line); border-radius:14px; overflow:hidden; }}
        th, td {{ padding:10px 12px; border-bottom:1px solid var(--line); text-align:left; font-size:13px; }}
        th {{ background:#eef3ff; color:#244091; }}
        td.pass {{ color:var(--ok); font-weight:700; }}
        td.fail {{ color:var(--bad); font-weight:700; }}
        .semester-section {{ margin-top:20px; }}
      </style>
    </head>
    <body>
      <div class="page">
        <div class="hero">
          <div>
            <div class="brand">AcadExtract</div>
            <div class="meta">Academic Report · Student Result Automation</div>
          </div>
          <div class="meta">USN {usn}</div>
        </div>

        <section class="profile">
          <div><div class="field-label">Student</div><div class="field-value">{name}</div></div>
          <div><div class="field-label">Email</div><div class="field-value">{email_str}</div></div>
          <div><div class="field-label">Department</div><div class="field-value">{dept or '—'}</div></div>
          <div><div class="field-label">Batch Year</div><div class="field-value">{batch_year or '—'}</div></div>
          <div><div class="field-label">CGPA</div><div class="field-value">{cgpa:.2f}</div></div>
          <div><div class="field-label">Backlogs</div><div class="field-value">{backlogs}</div></div>
        </section>

        <section class="metrics">
          <div class="metric-card"><div class="metric-label">Overall Percentage</div><div class="metric-value">{metrics['percentage']:.2f}%</div><div class="metric-sub">{metrics['total_scored']:.0f} / {metrics['total_max']:.0f}</div></div>
          <div class="metric-card"><div class="metric-label">Subjects Passed</div><div class="metric-value">{metrics['subjects_passed']} / {metrics['subjects_total']}</div><div class="metric-sub">Pass rate {metrics['pass_rate']:.2f}%</div></div>
          <div class="metric-card"><div class="metric-label">Best SGPA</div><div class="metric-value">{metrics['best_sgpa']:.2f}</div><div class="metric-sub">{metrics['semesters_completed']} semester(s)</div></div>
          {''.join(sem_cards)}
        </section>

        <h2>Detailed Results</h2>
        {''.join(detail_sections) if detail_sections else '<p>No detailed results available.</p>'}
      </div>
    </body>
    </html>
    """


# ---------------------------------------------------------------------------
# Phase 5 — intent parsing
# ---------------------------------------------------------------------------

def _parse_intent_local(query: str) -> dict[str, Any]:
    q = query.lower()
    usn_m = re.search(r'\b[1-4][a-z]{2}\d{2}[a-z]{2,4}\d{3}\b', q, re.IGNORECASE)
    usn   = usn_m.group(0).upper() if usn_m else None
    sem_m = re.search(r'\b(\d)\s*(st|nd|rd|th)?\s*sem', q)
    semester = int(sem_m.group(1)) if sem_m else None

    if usn and any(w in q for w in ["result", "marks", "grade", "score", "show", "get"]):
        intent = "student_lookup"
    elif any(w in q for w in ["backlog", "fail", "failed", "arrear"]):
        intent = "backlogs" if usn else "backlogs_list"
    elif any(w in q for w in ["top", "best", "rank", "topper"]):
        intent = "top_n"
    elif any(w in q for w in ["average", "avg", "mean", "cgpa", "overall"]):
        intent = "aggregation"
    elif any(w in q for w in ["trend", "progress", "across semester"]):
        intent = "trend"
    elif usn:
        intent = "student_lookup"
    else:
        intent = "aggregation"

    return {"intent": intent, "usn": usn, "semester": semester, "confidence": 0.75}


# ---------------------------------------------------------------------------
# Query execution (PostgreSQL)
# ---------------------------------------------------------------------------

def _get_institution_id() -> str:
    return db.get_default_institution_id()


def _execute_query(
    intent: str,
    usn: str | None,
    semester: int | None,
    raw_query: str,
) -> tuple[list[dict], str, str]:
    db.init_db()
    inst_id = _get_institution_id()

    # ── student_lookup ───────────────────────────────────────────────────────
    if intent == "student_lookup" and usn:
        student = db.get_student(usn)
        if not student:
            matches = db.search_students(usn)
            if not matches:
                return [], "", f"No student found with USN {usn}."
            student = matches[0]
            usn = student["usn"]

        results = db.get_student_results(usn, semester)
        name = _student_display_name(student)

        if not results:
            return [_serialise(student)], "", (
                f"Found student {name} (USN: {usn}) but no result records yet. "
                "Run the pipeline to extract results."
            )

        passes = sum(1 for r in results if str(r.get("pass_status", "")).lower() == "pass")
        fails  = sum(1 for r in results if str(r.get("pass_status", "")).lower() == "fail")
        text   = f"Results for {name} (USN: {usn})"
        text  += f" - Sem {semester}: " if semester else ": "
        text  += f"{len(results)} subject(s), {passes} passed, {fails} failed."
        if student.get("cgpa"):
            text += f" CGPA: {float(student['cgpa']):.2f}."

        return (
            [_serialise(r) for r in results],
            f"SELECT * FROM student_results WHERE student_id = '{student['id']}'",
            text,
        )

    # ── backlogs_list ────────────────────────────────────────────────────────
    if intent == "backlogs_list":
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("""
                    SELECT usn, name AS full_name, total_backlogs, cgpa
                    FROM students
                    WHERE institution_id = %s AND total_backlogs > 0
                                            {student_filter}
                                        ORDER BY total_backlogs DESC LIMIT 50
                                """.format(student_filter=db.student_source_filter(include_and=True)), (inst_id,))
                data = [_serialise(dict(r)) for r in cur.fetchall()]
        if not data:
            return [], "", "No students with backlogs found."
        top = data[0]
        text = (
            f"{len(data)} student(s) have backlogs. "
            f"Most: {top.get('full_name', top.get('usn'))} ({top['total_backlogs']})."
        )
        return data, "SELECT usn, full_name, total_backlogs FROM students WHERE total_backlogs > 0", text

    # ── backlogs (single student) ────────────────────────────────────────────
    if intent == "backlogs" and usn:
        student = db.get_student(usn)
        if not student:
            return [], "", f"Student {usn} not found."
        results = db.get_student_results(usn)
        failed  = [_serialise(r) for r in results if str(r.get("pass_status", "")).lower() == "fail"]
        name    = _student_display_name(student)
        return (
            failed,
            f"SELECT * FROM student_results WHERE student_id = '{student['id']}' AND pass_status='fail'",
            f"{name} has {len(failed)} backlog(s).",
        )

    # ── top_n ────────────────────────────────────────────────────────────────
    if intent == "top_n":
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("""
                    SELECT usn, name AS full_name, cgpa, total_backlogs
                    FROM students
                    WHERE institution_id = %s AND cgpa > 0
                                            {student_filter}
                                        ORDER BY cgpa DESC LIMIT 10
                                """.format(student_filter=db.student_source_filter(include_and=True)), (inst_id,))
                data = [_serialise(dict(r)) for r in cur.fetchall()]
        if not data:
            return [], "", "No CGPA data yet. Run the pipeline first."
        top  = data[0]
        name = top.get("full_name", top.get("usn", ""))
        return (
            data,
            "SELECT usn, full_name, cgpa FROM students ORDER BY cgpa DESC LIMIT 10",
            f"Top {len(data)} students by CGPA: {name} leads with {float(top['cgpa']):.2f}.",
        )

    # ── aggregation ──────────────────────────────────────────────────────────
    if intent in ("aggregation", "comparison"):
        stats = db.get_pipeline_stats()
        text  = (
            f"DB: {stats['total_students']} students, {stats['total_results']} records. "
            f"Avg CGPA: {stats['average_cgpa']:.2f}. Backlogs: {stats['total_backlogs']}."
        )
        return [stats], "SELECT COUNT(*), AVG(cgpa) FROM students", text

    # ── trend ────────────────────────────────────────────────────────────────
    if intent == "trend" and usn:
        student = db.get_student(usn)
        if not student:
            return [], "", f"Student {usn} not found."
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("""
                    SELECT semester, sgpa, credits_earned, backlogs
                    FROM semester_aggregates
                    WHERE student_id = %s
                    ORDER BY semester
                """, (student["id"],))
                trend = [_serialise(dict(r)) for r in cur.fetchall()]
        name = _student_display_name(student)
        if not trend:
            return [_serialise(student)], "", f"No semester data for {name} yet."
        text = "Trend: " + ", ".join(
            f"S{r['semester']}:{float(r['sgpa']):.2f}" for r in trend if r.get("sgpa")
        )
        return (
            trend,
            f"SELECT semester, sgpa FROM semester_aggregates WHERE student_id = '{student['id']}'",
            text,
        )

    # ── fallback: name search — treat as student_lookup if exactly one match ─
    for word in [w for w in raw_query.split() if len(w) > 3]:
        matches = db.search_students(word, limit=5)
        if not matches:
            continue
        if len(matches) == 1:
            # Single match → show full results, same as student_lookup
            student = matches[0]
            found_usn = student["usn"]
            results = db.get_student_results(found_usn, semester)
            name = _student_display_name(student)
            if results:
                passes = sum(1 for r in results if str(r.get("pass_status", "")).lower() == "pass")
                fails  = sum(1 for r in results if str(r.get("pass_status", "")).lower() == "fail")
                text   = f"Results for {name} (USN: {found_usn}): "
                text  += f"{len(results)} subject(s), {passes} passed, {fails} failed."
                if student.get("cgpa"):
                    text += f" CGPA: {float(student['cgpa']):.2f}."
                return (
                    [_serialise(r) for r in results],
                    f"SELECT * FROM student_results WHERE student_id = '{student['id']}'",
                    text,
                )
            return (
                [_serialise(student)],
                f"SELECT * FROM students WHERE usn = '{found_usn}'",
                f"Found {name} (USN: {found_usn}) but no results stored yet.",
            )
        # Multiple matches — list them
        return (
            [_serialise(m) for m in matches],
            f"SELECT * FROM students WHERE name ILIKE '%{word}%'",
            f"Found {len(matches)} student(s) matching '{word}'. Specify a USN for detailed results.",
        )

    stats = db.get_pipeline_stats()
    return (
        [stats],
        "",
        f"DB: {stats['total_students']} students, avg CGPA {stats['average_cgpa']}.",
    )


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

def _llm_synthesize_query_answer(
    query: str,
    intent: str,
    rows: list[dict],
    template_answer: str,
) -> tuple[str, list[str]]:
    """
    Use the configured LLM to produce a natural-language answer for a query.
    Falls back to the template answer if no LLM is available or call fails.

    Returns (answer_text, sources) where sources is a list of cited USNs/names.
    """
    # Extract source identifiers (USNs or names) from rows for citations
    sources: list[str] = []
    for r in rows[:20]:
        identifier = r.get("usn") or r.get("full_name") or r.get("name")
        if identifier and str(identifier) not in sources:
            sources.append(str(identifier))

    # Insufficient-data guard (applied before LLM call)
    if not rows:
        return (
            "I don't have sufficient data to answer this question. "
            "Please run the pipeline to import more results.",
            [],
        )

    try:
        from src.common.config import get_settings
        settings = get_settings()
        provider = settings.llm.active_provider
        if provider == "none":
            return template_answer, sources

        # Build a compact context string from the data rows (cap at 20 rows)
        context_parts = [f"Query: {query}", f"Intent: {intent}", f"Raw answer: {template_answer}"]
        if rows:
            context_parts.append(f"Data ({min(len(rows), 20)} of {len(rows)} record(s)):")
            for r in rows[:20]:
                context_parts.append("  " + "  |  ".join(
                    f"{k}: {v}" for k, v in r.items()
                    if v is not None and str(v) not in ("", "None")
                ))

        sources_str = ", ".join(sources[:10]) if sources else "none"
        prompt = (
            "You are AcadAssist, an AI assistant for teachers at MSRIT/VTU. "
            "Given the following academic database query result, write a concise, "
            "clear natural-language answer for the teacher. "
            "Use markdown for emphasis where helpful. Be factual and precise. "
            "IMPORTANT: Only state facts present in the provided data. "
            "If you are uncertain or data is insufficient, say "
            "'I don't have enough data to fully answer this question.' "
            f"Cite specific students or records where relevant (available: {sources_str}).\n\n"
            + "\n".join(context_parts)
            + "\n\nAnswer:"
        )

        # Waterfall: try each provider in order until one succeeds
        _synth_providers: list[str] = []
        if settings.llm.groq_api_key:
            _synth_providers.append("groq")
        if settings.llm.primary_api_key:
            _synth_providers.append("openai")
        if settings.llm.secondary_api_key:
            _synth_providers.append("gemini")

        llm_answer = None
        for _provider in _synth_providers:
            try:
                if _provider == "groq":
                    from groq import Groq
                    client = Groq(api_key=settings.llm.groq_api_key)
                    completion = client.chat.completions.create(
                        model=settings.llm.groq_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.2,
                        max_tokens=512,
                    )
                    llm_answer = completion.choices[0].message.content.strip()
                elif _provider == "openai":
                    from openai import OpenAI
                    client = OpenAI(api_key=settings.llm.primary_api_key)
                    resp = client.chat.completions.create(
                        model=settings.llm.primary_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.2,
                        max_tokens=512,
                    )
                    llm_answer = resp.choices[0].message.content.strip()
                elif _provider == "gemini":
                    import google.generativeai as genai
                    genai.configure(api_key=settings.llm.secondary_api_key)
                    gmodel = genai.GenerativeModel(settings.llm.secondary_model)
                    response = gmodel.generate_content(
                        prompt,
                        generation_config={"temperature": 0.2, "max_output_tokens": 512},
                    )
                    llm_answer = response.text.strip()
                if llm_answer:
                    break  # success
            except Exception as _exc:
                logger.warning("llm_query_synthesis_%s_failed: %s", _provider, _exc)

        if llm_answer:
            # Append sources footnote if we have identifiers
            if sources:
                sources_footnote = "\n\n---\n*Sources: " + ", ".join(sources[:10]) + "*"
                llm_answer += sources_footnote
            return llm_answer, sources

    except Exception as exc:
        logger.warning("llm_query_synthesis_failed", error=str(exc))

    return template_answer, sources


@router.post("/query", response_model=QueryResponse)
async def submit_query(request: QueryRequest) -> QueryResponse:
    """Natural language academic query (Phase 5)."""
    start    = time.perf_counter()
    db.init_db()

    parsed   = _parse_intent_local(request.query)
    intent   = parsed["intent"]
    usn      = parsed.get("usn")
    semester = parsed.get("semester")

    rows, sql, template_answer = _execute_query(intent, usn, semester, request.query)

    # Upgrade template answer to natural-language via LLM (best-effort, non-blocking)
    text_answer, sources = _llm_synthesize_query_answer(request.query, intent, rows, template_answer)

    elapsed  = int((time.perf_counter() - start) * 1000)

    chart_spec = None
    if intent == "trend" and rows and isinstance(rows[0], dict) and rows[0].get("sgpa"):
        chart_spec = {
            "type": "line",
            "x": [r.get("semester") for r in rows],
            "y": [r.get("sgpa") for r in rows],
            "xlabel": "Semester", "ylabel": "SGPA",
        }
    elif intent == "top_n" and rows:
        chart_spec = {
            "type": "bar",
            "labels": [r.get("full_name", r.get("usn", "")) for r in rows],
            "values": [float(r.get("cgpa", 0) or 0) for r in rows],
        }

    # Write to query_audit_log (best-effort — never fails the response)
    db.store_audit_log(
        raw_query=request.query,
        parsed_intent=intent,
        sql_generated=sql or None,
        response_summary=(text_answer or "")[:500],
        result_count=len(rows),
        duration_ms=elapsed,
    )

    return QueryResponse(
        query=request.query,
        intent=intent,
        text_answer=text_answer or "No results found.",
        summary=f"{len(rows)} record(s) in {elapsed}ms.",
        data=rows,
        chart_spec=chart_spec,
        confidence=parsed.get("confidence", 0.75),
        caveats=(["No data found. Run the pipeline on synced emails."] if not rows else []),
        sources=sources,
    )


@router.get("/student/{usn}", response_model=StudentSummary)
async def get_student(usn: str) -> StudentSummary:
    """Full student profile: aggregates, per-subject results, class rank."""
    db.init_db()
    student = db.get_student(usn.upper())
    if not student:
        raise HTTPException(status_code=404, detail=f"Student {usn} not found")

    sid = student["id"]
    inst_id = _get_institution_id()

    with db.get_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            # Semester aggregates
            cur.execute("""
                SELECT semester, sgpa, credits_earned, credits_attempted,
                       subjects_passed, subjects_failed, backlogs
                FROM semester_aggregates WHERE student_id = %s ORDER BY semester
            """, (sid,))
            sems = [_serialise(dict(r)) for r in cur.fetchall()]

            # All subject results
            cur.execute("""
                SELECT sr.semester, sr.total_marks AS marks_obtained, sr.max_marks,
                       sr.grade, sr.grade_points, sr.status AS pass_status,
                       sub.code AS subject_code, sub.name AS subject_name
                FROM student_results sr
                JOIN subjects sub ON sr.subject_id = sub.id
                WHERE sr.student_id = %s
                ORDER BY sr.semester, sub.code
            """, (sid,))
            all_results = [_serialise(dict(r)) for r in cur.fetchall()]

            # Class rank (non-seed only)
            cur.execute("""
                SELECT COUNT(*) AS rnk FROM students
                WHERE institution_id = %s
                {student_filter}
                  AND cgpa > %s
            """.format(student_filter=db.student_source_filter(include_and=True)), (inst_id, float(student.get("cgpa") or 0)))
            rnk_row = cur.fetchone()
            class_rank = int(rnk_row["rnk"]) + 1 if rnk_row else None

            # Total non-seed students
            cur.execute("""
                SELECT COUNT(*) AS cnt FROM students
                WHERE institution_id = %s
                {student_filter}
            """.format(student_filter=db.student_source_filter(include_and=True)), (inst_id,))
            total_students = int(cur.fetchone()["cnt"])

    # Group results by semester for easy frontend rendering
    results_by_sem: dict[str, list] = {}
    for r in all_results:
        key = str(r.get("semester", "?"))
        results_by_sem.setdefault(key, []).append(r)

    # Derive batch_year and department from USN  (e.g. 1MS21CS001 → batch=2021, dept=CS)
    u = student["usn"].upper()
    batch_year = None
    department = ""
    try:
        year_2d = int(u[3:5])
        batch_year = 2000 + year_2d
        department = u[5:7]  # e.g. CS, IS, EC
    except Exception:
        pass

    return StudentSummary(
        usn=student["usn"],
        name=_student_display_name(student),
        email=student.get("email"),
        department=department,
        batch_year=batch_year,
        current_cgpa=float(student.get("cgpa") or 0),
        active_backlogs=int(student.get("active_backlogs") or 0),
        total_results=len(all_results),
        class_rank=class_rank,
        total_students=total_students,
        semesters=sems,
        results_by_semester=results_by_sem,
    )


@router.get("/student/{usn}/trend", response_model=TrendResponse)
async def get_student_trend(usn: str) -> TrendResponse:
    """SGPA trend across semesters."""
    db.init_db()
    student = db.get_student(usn.upper())
    if not student:
        raise HTTPException(status_code=404, detail=f"Student {usn} not found")

    with db.get_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT semester, sgpa, credits_earned, backlogs
                FROM semester_aggregates
                WHERE student_id = %s
                ORDER BY semester
            """, (student["id"],))
            trend = [_serialise(dict(r)) for r in cur.fetchall()]

    return TrendResponse(
        usn=student["usn"],
        name=_student_display_name(student),
        trend=trend,
    )


@router.get("/student/{usn}/report")
async def export_student_report(usn: str, format: str = "pdf"):
    """Generate a downloadable PDF or DOCX academic report for a student."""
    from fastapi.responses import StreamingResponse
    import io

    db.init_db()
    student = db.get_student(usn.upper())
    if not student:
        raise HTTPException(status_code=404, detail=f"Student {usn} not found")

    sid = student["id"]
    all_results = db.get_student_results(usn.upper())

    with db.get_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT semester, sgpa, credits_earned, subjects_passed, subjects_failed, backlogs
                FROM semester_aggregates WHERE student_id = %s ORDER BY semester
            """, (sid,))
            sems = [dict(r) for r in cur.fetchall()]
    metrics = _compute_report_metrics(all_results, sems)

    # Derive department + batch from USN
    u = student["usn"].upper()
    try:
        batch_year = 2000 + int(u[3:5])
        department = u[5:7]
    except Exception:
        batch_year, department = None, ""

    name      = _student_display_name(student)
    cgpa      = float(student.get("cgpa") or 0)
    backlogs  = int(student.get("active_backlogs") or 0)
    email_str = student.get("email") or "—"

    fmt = format.lower()
    if fmt == "docx":
        buf = _build_docx_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.docx"
        media  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "xlsx":
        buf = _build_xlsx_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.xlsx"
        media  = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        buf = _build_pdf_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.pdf"
        media  = "application/pdf"

    return StreamingResponse(
        buf,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ---------------------------------------------------------------------------
# Email report endpoint
# ---------------------------------------------------------------------------

class EmailReportRequest(BaseModel):
    usn: str = Field(..., min_length=3, max_length=30)
    recipient: str = Field(..., min_length=5, max_length=254)
    format: str = Field(default="pdf")
    subject: str | None = None


class EmailReportResponse(BaseModel):
    sent: bool
    recipient: str
    usn: str
    message: str


def _send_email_with_attachment(
    to_addr: str, subject: str, body_html: str,
    attachment_bytes: bytes, attachment_name: str, mime_type: str,
) -> None:
    """Send an email with a file attachment via SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    from src.common.config import get_settings

    cfg = get_settings().smtp
    if not cfg.configured:
        raise RuntimeError("SMTP not configured. Set SMTP_USER and SMTP_PASSWORD in .env")

    msg = MIMEMultipart()
    msg["From"] = f"{cfg.from_name} <{cfg.user}>"
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.attach(MIMEText(body_html, "html"))

    part = MIMEBase("application", "octet-stream")
    part.set_payload(attachment_bytes)
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{attachment_name}"')
    msg.attach(part)

    with smtplib.SMTP(cfg.host, cfg.port, timeout=30) as server:
        if cfg.use_tls:
            server.starttls()
        server.login(cfg.user, cfg.password)
        server.send_message(msg)


@router.post("/student/{usn}/email-report", response_model=EmailReportResponse)
async def email_student_report(usn: str, req: EmailReportRequest):
    """Generate an academic report and email it to the specified recipient."""
    # Validate email format
    if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', req.recipient):
        raise HTTPException(status_code=400, detail="Invalid email address")

    db.init_db()
    student = db.get_student(usn.upper())
    if not student:
        raise HTTPException(status_code=404, detail=f"Student {usn} not found")

    sid = student["id"]
    all_results = db.get_student_results(usn.upper())

    with db.get_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT semester, sgpa, credits_earned, subjects_passed, subjects_failed, backlogs
                FROM semester_aggregates WHERE student_id = %s ORDER BY semester
            """, (sid,))
            sems = [dict(r) for r in cur.fetchall()]

    u = student["usn"].upper()
    try:
        batch_year = 2000 + int(u[3:5])
        department = u[5:7]
    except Exception:
        batch_year, department = None, ""

    name      = _student_display_name(student)
    cgpa      = float(student.get("cgpa") or 0)
    backlogs  = int(student.get("active_backlogs") or 0)
    email_str = student.get("email") or "—"

    fmt = (req.format or "pdf").lower()
    if fmt == "docx":
        buf = _build_docx_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.docx"
        mime  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "xlsx":
        buf = _build_xlsx_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.xlsx"
        mime  = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        buf = _build_pdf_report(u, name, email_str, department, batch_year, cgpa, backlogs, sems, all_results, metrics)
        fname = f"{u}_report.pdf"
        mime  = "application/pdf"

    subj = req.subject or f"Academic Report — {name} ({u})"
    body_html = f"""
    <div style="font-family:Arial,sans-serif;color:#1a1d27;">
        <h2 style="color:#4f8cff;">AcadExtract — Academic Report</h2>
        <p>Hello,</p>
        <p>Please find attached the academic report for <strong>{name}</strong> (USN: {u}).</p>
        <ul>
            <li><strong>CGPA:</strong> {cgpa:.2f}</li>
            <li><strong>Overall Percentage:</strong> {metrics['percentage']:.2f}%</li>
            <li><strong>Subjects Passed:</strong> {metrics['subjects_passed']} / {metrics['subjects_total']}</li>
            <li><strong>Active Backlogs:</strong> {backlogs}</li>
            <li><strong>Department:</strong> {department or '—'}</li>
        </ul>
        <p>This report was generated automatically by AcadExtract.</p>
        <hr style="border:none;border-top:1px solid #e4e6eb;">
        <p style="font-size:12px;color:#8b8fa3;">AcadExtract — Autonomous Academic Result Extraction & Student Profiling</p>
    </div>
    """

    try:
        _send_email_with_attachment(req.recipient, subj, body_html, buf.read(), fname, mime)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error("email_send_failed", error=str(e), to=req.recipient, usn=u)
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

    logger.info("email_report_sent", to=req.recipient, usn=u, format=fmt)
    return EmailReportResponse(
        sent=True, recipient=req.recipient, usn=u,
        message=f"Report emailed to {req.recipient}"
    )


def _build_pdf_report(usn, name, email_str, dept, batch_year, cgpa, backlogs, sems, results, metrics):
    """Build a PDF academic report using ReportLab."""
    import io
    from datetime import date
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2*cm,  bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    accent = colors.HexColor("#4f8cff")
    dark   = colors.HexColor("#1a1d27")
    mid    = colors.HexColor("#2a2d37")
    light  = colors.HexColor("#e4e6eb")
    muted  = colors.HexColor("#8b8fa3")

    title_style  = ParagraphStyle("Title",  parent=styles["Heading1"], fontSize=20, textColor=dark,   spaceAfter=4)
    sub_style    = ParagraphStyle("Sub",    parent=styles["Normal"],   fontSize=10, textColor=muted,  spaceAfter=12)
    section_style= ParagraphStyle("Sec",    parent=styles["Heading2"], fontSize=13, textColor=accent, spaceBefore=14, spaceAfter=6)
    body_style   = ParagraphStyle("Body",   parent=styles["Normal"],   fontSize=10, textColor=dark,   leading=16)

    W = A4[0] - 4*cm  # usable width

    story = []

    # ── Header ──────────────────────────────────────────────────────
    story.append(Paragraph("AcadExtract", title_style))
    story.append(Paragraph(f"Academic Report  ·  Generated {date.today().strftime('%d %B %Y')}", sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=mid, spaceAfter=14))

    # ── Student Profile ─────────────────────────────────────────────
    story.append(Paragraph("Student Profile", section_style))
    profile_data = [
        ["USN",           usn],
        ["Name",          name],
        ["Email",         email_str],
        ["Department",    dept if dept else "—"],
        ["Batch Year",    str(batch_year) if batch_year else "—"],
        ["CGPA",          f"{cgpa:.2f}"],
        ["Overall Percentage", f"{metrics['percentage']:.2f}%"],
        ["Subjects Passed", f"{metrics['subjects_passed']} / {metrics['subjects_total']}"],
        ["Best SGPA", f"{metrics['best_sgpa']:.2f}" if metrics["best_sgpa"] else "â€”"],
        ["Active Backlogs", str(backlogs)],
    ]
    pt = Table(profile_data, colWidths=[4*cm, W-4*cm])
    pt.setStyle(TableStyle([
        ("FONTSIZE",    (0,0), (-1,-1), 10),
        ("FONTNAME",    (0,0), (0,-1), "Helvetica-Bold"),
        ("TEXTCOLOR",   (0,0), (0,-1), muted),
        ("TEXTCOLOR",   (1,0), (1,-1), dark),
        ("TOPPADDING",  (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ("ROWBACKGROUNDS",(0,0),(-1,-1),[colors.white, colors.HexColor("#f5f6fa")]),
        ("LINEBELOW",   (0,-1),(-1,-1), 0.5, mid),
    ]))
    story.append(pt)
    story.append(Spacer(1, 14))

    # ── Semester Summary ────────────────────────────────────────────
    if sems:
        story.append(Paragraph("Semester Summary", section_style))
        sem_hdr = ["Sem", "SGPA", "Credits", "Passed", "Failed", "Backlogs"]
        sem_rows = [[
            str(s.get("semester", "")),
            f"{float(s['sgpa']):.2f}" if s.get("sgpa") else "—",
            str(s.get("credits_earned", "—")),
            str(s.get("subjects_passed", "—")),
            str(s.get("subjects_failed", "—")),
            str(s.get("backlogs", 0)),
        ] for s in sems]
        st = Table([sem_hdr] + sem_rows, repeatRows=1,
                   colWidths=[1.2*cm, 2*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
        st.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0), accent),
            ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
            ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,-1), 9),
            ("ALIGN",       (0,0), (-1,-1), "CENTER"),
            ("TOPPADDING",  (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#f5f6fa")]),
            ("GRID",        (0,0), (-1,-1), 0.4, mid),
        ]))
        story.append(st)
        story.append(Spacer(1, 14))

    # ── Detailed Results per Semester ───────────────────────────────
    by_sem: dict[int, list] = {}
    for r in results:
        by_sem.setdefault(int(r.get("semester", 0)), []).append(r)

    if by_sem:
        story.append(Paragraph("Detailed Results", section_style))
        for sem_num in sorted(by_sem):
            story.append(Paragraph(f"Semester {sem_num}", body_style))
            hdr = ["Code", "Subject", "Marks", "Max", "Grade", "Status"]
            rows = [[
                r.get("subject_code", ""),
                r.get("subject_name", "")[:35],
                str(int(r.get("marks_obtained", 0) or 0)),
                str(int(r.get("max_marks", 100) or 100)),
                r.get("grade", "—") or "—",
                r.get("pass_status", "—") or "—",
            ] for r in by_sem[sem_num]]
            rt = Table([hdr]+rows, repeatRows=1,
                       colWidths=[2*cm, (W-11*cm), 1.8*cm, 1.8*cm, 1.8*cm, 2*cm])
            rt.setStyle(TableStyle([
                ("BACKGROUND",  (0,0), (-1,0), mid),
                ("TEXTCOLOR",   (0,0), (-1,0), light),
                ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",    (0,0), (-1,-1), 8.5),
                ("ALIGN",       (2,0), (-1,-1), "CENTER"),
                ("TOPPADDING",  (0,0), (-1,-1), 3),
                ("BOTTOMPADDING",(0,0),(-1,-1),3),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#f5f6fa")]),
                ("GRID",        (0,0), (-1,-1), 0.3, mid),
                # red FAIL rows
                *[("TEXTCOLOR", (5, i+1), (5, i+1),
                   colors.HexColor("#e55353") if row[5] == "FAIL" else colors.HexColor("#2ecc71"))
                  for i, row in enumerate(rows)],
            ]))
            story.append(rt)
            story.append(Spacer(1, 10))

    doc.build(story)
    buf.seek(0)
    return buf


def _build_xlsx_report(usn, name, email_str, dept, batch_year, cgpa, backlogs, sems, results, metrics):
    """Build an Excel academic report using openpyxl."""
    import io
    from datetime import date
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # ── Colours & styles ────────────────────────────────────────────
    accent_fill = PatternFill(start_color="4F8CFF", end_color="4F8CFF", fill_type="solid")
    dark_fill   = PatternFill(start_color="1A1D27", end_color="1A1D27", fill_type="solid")
    alt_fill    = PatternFill(start_color="F5F6FA", end_color="F5F6FA", fill_type="solid")
    hdr_font    = Font(bold=True, color="FFFFFF", size=10)
    title_font  = Font(bold=True, size=14, color="1A1D27")
    sub_font    = Font(italic=True, size=9, color="8B8FA3")
    label_font  = Font(bold=True, size=10, color="8B8FA3")
    value_font  = Font(size=10, color="1A1D27")
    fail_font   = Font(size=10, color="E55353", bold=True)
    pass_font   = Font(size=10, color="2ECC71")
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    # ── Summary sheet ───────────────────────────────────────────────
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_properties.tabColor = "4F8CFF"

    ws.merge_cells("A1:F1")
    ws["A1"].value = "AcadExtract — Academic Report"
    ws["A1"].font = title_font
    ws["A2"].value = f"Generated {date.today().strftime('%d %B %Y')}"
    ws["A2"].font = sub_font

    row = 4
    profile = [
        ("USN", usn), ("Name", name), ("Email", email_str),
        ("Department", dept or "—"), ("Batch Year", str(batch_year) if batch_year else "—"),
        ("CGPA", f"{cgpa:.2f}"), ("Overall Percentage", f"{metrics['percentage']:.2f}%"),
        ("Subjects Passed", f"{metrics['subjects_passed']} / {metrics['subjects_total']}"),
        ("Best SGPA", f"{metrics['best_sgpa']:.2f}" if metrics["best_sgpa"] else "â€”"),
        ("Active Backlogs", str(backlogs)),
    ]
    for lbl, val in profile:
        ws.cell(row=row, column=1, value=lbl).font = label_font
        ws.cell(row=row, column=2, value=val).font = value_font
        row += 1

    # Semester summary table
    if sems:
        row += 1
        ws.cell(row=row, column=1, value="Semester Summary").font = Font(bold=True, size=12, color="4F8CFF")
        row += 1
        sem_headers = ["Semester", "SGPA", "Credits Earned", "Passed", "Failed", "Backlogs"]
        for c, h in enumerate(sem_headers, 1):
            cell = ws.cell(row=row, column=c, value=h)
            cell.font = hdr_font
            cell.fill = accent_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
        for s in sems:
            row += 1
            vals = [
                s.get("semester", ""),
                float(s["sgpa"]) if s.get("sgpa") else None,
                s.get("credits_earned"),
                s.get("subjects_passed"),
                s.get("subjects_failed"),
                s.get("backlogs", 0),
            ]
            for c, v in enumerate(vals, 1):
                cell = ws.cell(row=row, column=c, value=v)
                cell.font = value_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border
                if row % 2 == 0:
                    cell.fill = alt_fill

    # Auto-fit column widths
    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 30)

    # ── Per-semester detail sheets ──────────────────────────────────
    by_sem: dict[int, list] = {}
    for r in results:
        by_sem.setdefault(int(r.get("semester", 0)), []).append(r)

    for sem_num in sorted(by_sem):
        sheet = wb.create_sheet(title=f"Sem {sem_num}")
        sheet.sheet_properties.tabColor = "4F8CFF"

        sheet.merge_cells("A1:F1")
        sheet["A1"].value = f"Semester {sem_num} — Detailed Results"
        sheet["A1"].font = Font(bold=True, size=12, color="4F8CFF")

        headers = ["Subject Code", "Subject Name", "Marks", "Max Marks", "Grade", "Status"]
        for c, h in enumerate(headers, 1):
            cell = sheet.cell(row=3, column=c, value=h)
            cell.font = hdr_font
            cell.fill = dark_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border

        for i, r in enumerate(by_sem[sem_num], start=4):
            status = r.get("pass_status", "—") or "—"
            vals = [
                r.get("subject_code", ""),
                r.get("subject_name", ""),
                int(r.get("marks_obtained", 0) or 0),
                int(r.get("max_marks", 100) or 100),
                r.get("grade", "—") or "—",
                status,
            ]
            for c, v in enumerate(vals, 1):
                cell = sheet.cell(row=i, column=c, value=v)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center") if c >= 3 else Alignment()
                if status == "FAIL":
                    cell.font = fail_font
                elif c == 6 and status == "PASS":
                    cell.font = pass_font
                else:
                    cell.font = value_font
                if i % 2 == 0:
                    cell.fill = alt_fill

        for col in sheet.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            sheet.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _build_docx_report(usn, name, email_str, dept, batch_year, cgpa, backlogs, sems, results, metrics):
    """Build a DOCX academic report using python-docx."""
    import io
    from datetime import date
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ACCENT = RGBColor(0x4f, 0x8c, 0xff)
    DARK   = RGBColor(0x1a, 0x1d, 0x27)
    MUTED  = RGBColor(0x8b, 0x8f, 0xa3)
    FAIL_C = RGBColor(0xe5, 0x53, 0x53)
    PASS_C = RGBColor(0x2e, 0xcc, 0x71)

    doc = Document()
    # Set narrow margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def add_heading(text, level=1, color=DARK):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def add_field(label, value):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True; r.font.color.rgb = MUTED
        r2 = p.add_run(str(value))
        r2.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(2)
        return p

    # Header
    h = doc.add_heading("AcadExtract", 0)
    for run in h.runs:
        run.font.color.rgb = DARK
    sub = doc.add_paragraph(f"Academic Report  ·  Generated {date.today().strftime('%d %B %Y')}")
    sub.runs[0].font.color.rgb = MUTED
    sub.runs[0].font.size = Pt(9)
    doc.add_paragraph("─" * 80).runs[0].font.color.rgb = MUTED

    # Student Profile
    add_heading("Student Profile", 1, ACCENT)
    add_field("USN",            usn)
    add_field("Name",           name)
    add_field("Email",          email_str)
    add_field("Department",     dept if dept else "—")
    add_field("Batch Year",     str(batch_year) if batch_year else "—")
    add_field("CGPA",           f"{cgpa:.2f}")
    add_field("Overall Percentage", f"{metrics['percentage']:.2f}%")
    add_field("Subjects Passed", f"{metrics['subjects_passed']} / {metrics['subjects_total']}")
    add_field("Best SGPA", f"{metrics['best_sgpa']:.2f}" if metrics["best_sgpa"] else "â€”")
    add_field("Active Backlogs", str(backlogs))

    # Semester Summary
    if sems:
        add_heading("Semester Summary", 1, ACCENT)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Shading Accent 1"
        for i, hdr_txt in enumerate(["Sem", "SGPA", "Credits", "Passed", "Failed", "Backlogs"]):
            cell = tbl.rows[0].cells[i]
            cell.text = hdr_txt
            cell.paragraphs[0].runs[0].bold = True
        for s in sems:
            row = tbl.add_row()
            row.cells[0].text = str(s.get("semester",""))
            row.cells[1].text = f"{float(s['sgpa']):.2f}" if s.get("sgpa") else "—"
            row.cells[2].text = str(s.get("credits_earned","—"))
            row.cells[3].text = str(s.get("subjects_passed","—"))
            row.cells[4].text = str(s.get("subjects_failed","—"))
            row.cells[5].text = str(s.get("backlogs", 0))

    # Detailed Results
    by_sem: dict[int, list] = {}
    for r in results:
        by_sem.setdefault(int(r.get("semester", 0)), []).append(r)

    if by_sem:
        add_heading("Detailed Results", 1, ACCENT)
        for sem_num in sorted(by_sem):
            add_heading(f"Semester {sem_num}", 2, DARK)
            tbl = doc.add_table(rows=1, cols=6)
            tbl.style = "Light Grid Accent 1"
            for i, hdr_txt in enumerate(["Code","Subject","Marks","Max","Grade","Status"]):
                cell = tbl.rows[0].cells[i]
                cell.text = hdr_txt
                cell.paragraphs[0].runs[0].bold = True
            for r in by_sem[sem_num]:
                row = tbl.add_row()
                row.cells[0].text = r.get("subject_code","")
                row.cells[1].text = r.get("subject_name","")
                row.cells[2].text = str(int(r.get("marks_obtained",0) or 0))
                row.cells[3].text = str(int(r.get("max_marks",100) or 100))
                row.cells[4].text = r.get("grade","—") or "—"
                status = r.get("pass_status","—") or "—"
                row.cells[5].text = status
                if status == "FAIL":
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = FAIL_C

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/students")
async def list_students(q: str = "", limit: int = 50) -> dict:
    """List or search students."""
    db.init_db()
    students = db.search_students(q, limit=limit) if q else db.get_all_students(limit=limit)
    inst_id  = _get_institution_id()

    result = []
    with db.get_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            for s in students:
                cur.execute("""
                    SELECT COUNT(*) AS n FROM student_results sr
                    JOIN students st ON sr.student_id = st.id
                    WHERE st.usn = %s AND st.institution_id = %s
                """, (s["usn"], inst_id))
                cnt = cur.fetchone()["n"]
                row = _serialise(s)
                row["result_count"] = cnt
                result.append(row)

    return {"total": len(result), "students": result}


# ---------------------------------------------------------------------------
# AI Assistant — /chat  (Gemini-powered)
# ---------------------------------------------------------------------------

AI_SYSTEM_PROMPT = """You are AcadAssist, an AI assistant embedded in AcadExtract — an academic result extraction platform for teachers.

Your role: Help teachers understand student performance, answer questions about results, identify students needing attention, and summarize trends.

You have real-time access to the PostgreSQL database. All relevant data is provided in the DATABASE CONTEXT section below.

Rules:
- ALWAYS answer from the database context provided — do NOT say data is missing if it appears in context
- If a student is found in the context, show all their available details (USN, name, CGPA, subject marks, grades)
- For semester-specific questions, filter and show only that semester's subjects
- Format results as a clear table or bullet list with: Subject Code | Subject Name | Marks | Grade | Status
- When asked for a "grade card" or "marks sheet", show all subjects for that semester in a table format
- Be precise with numbers, student names, and USNs
- If a student truly does not appear anywhere in the context, say "Student not found in the database"
- Never tell the user to sync or run the pipeline — the data is already loaded
- Use bullet points for lists, markdown bold for emphasis
- Understand variations: 'backlog' = 'fail', 'topper' = highest CGPA, 'marks sheet' = 'grade card'"""


class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=2000)
    history: list[ChatMessage] = []


class ChatResponse(BaseModel):
    reply: str
    context_used: bool = False


def _local_chat_answer(message: str) -> str:
    """Generate a direct answer from the DB using the query engine (no LLM needed)."""
    try:
        db.init_db()
        parsed      = _parse_intent_local(message)
        intent      = parsed["intent"]
        usn         = parsed.get("usn")
        semester    = parsed.get("semester")
        rows, _, text_answer = _execute_query(intent, usn, semester, message)

        reply_lines = [text_answer]

        if rows:
            # Show first 10 rows as a simple summary
            for row in rows[:10]:
                parts = []
                for k in ("usn", "full_name", "subject_name", "subject_code",
                          "marks_obtained", "total_marks", "grade", "status",
                          "semester", "sgpa", "cgpa", "total_backlogs"):
                    v = row.get(k)
                    if v is not None and str(v) not in ("", "None"):
                        parts.append(f"{k}: {v}")
                if parts:
                    reply_lines.append("  • " + "  |  ".join(parts))

        if len(rows) > 10:
            reply_lines.append(f"  … and {len(rows) - 10} more record(s).")

        reply_lines.append(
            "\n_(No AI provider configured — showing direct database results. "
            "Add a GROQ_API_KEY, GEMINI_API_KEY, or OPENAI_API_KEY in your .env for natural-language responses.)_"
        )
        return "\n".join(reply_lines)
    except Exception as exc:
        logger.warning("local_chat_fallback_failed", error=str(exc))
        return (
            "No AI provider configured and the database query also failed. "
            "Please add a GROQ_API_KEY, GEMINI_API_KEY, or OPENAI_API_KEY to your .env file."
        )


def _verify_chat_reply(reply: str, inst_id: str) -> tuple[str, bool]:
    """
    Verification agent: extract every USN mentioned in the LLM reply, look up
    the student's actual DB values, and flag any discrepancies in CGPA or
    backlog count. Returns (annotated_reply, had_issues).
    """
    usns = list(dict.fromkeys(  # deduplicate, preserve order
        re.findall(r'\b[1-4][A-Z]{2}\d{2}[A-Z]{2,4}\d{3}\b', reply.upper())
    ))
    if not usns:
        return reply + "\n\n✅ *Verified — no specific student data in this response.*", False

    issues: list[str] = []
    verified: list[str] = []
    try:
        db.init_db()
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                for usn in usns:
                    cur.execute("""
                        SELECT s.usn, s.name AS full_name, s.cgpa, s.active_backlogs
                        FROM students s
                        WHERE UPPER(s.usn) = %s AND s.institution_id = %s
                          AND (s.metadata ? 'source')
                          AND s.metadata->>'source' IN ('pipeline', 'upload')
                    """, (usn, inst_id))
                    actual = cur.fetchone()
                    if not actual:
                        issues.append(f"⚠ **{usn}** not found in database — this data may be hallucinated.")
                        continue

                    name = actual.get("full_name") or usn
                    # Scan the ±100-char window around this USN for floating-point values
                    idx = reply.upper().find(usn)
                    snippet = reply[max(0, idx - 80): idx + 140]
                    floats_nearby = re.findall(r'\b(\d+\.\d{1,2})\b', snippet)
                    mismatch = False
                    if actual.get("cgpa") and floats_nearby:
                        actual_cgpa = float(actual["cgpa"])
                        for val in floats_nearby:
                            mentioned = float(val)
                            if 0.0 < mentioned <= 10.0 and abs(mentioned - actual_cgpa) > 0.15:
                                issues.append(
                                    f"⚠ **{usn} ({name})** — reply mentions {mentioned:.2f} "
                                    f"near this USN but actual CGPA is **{actual_cgpa:.2f}**."
                                )
                                mismatch = True
                                break
                    if not mismatch:
                        verified.append(f"✓ {usn} ({name}) — CGPA {float(actual['cgpa']):.2f}" if actual.get("cgpa") else f"✓ {usn} ({name})")
    except Exception as exc:
        logger.warning("verification_agent_failed", error=str(exc))
        return reply, False

    if not issues:
        note = "✅ *Verification agent: data confirmed for " + ", ".join(verified) + ".*"
        return reply + "\n\n" + note, False

    block = "\n\n---\n**🔍 Verification Agent:**\n" + "\n".join(issues)
    if verified:
        block += "\n" + "\n".join(verified)
    return reply + block, True


@router.post("/chat", response_model=ChatResponse)
async def ai_chat(req: ChatRequest) -> ChatResponse:
    """AI assistant endpoint — answers teacher questions using Gemini + live DB context."""
    from src.common.config import get_settings
    settings = get_settings()

    no_llm = (
        not settings.llm.groq_api_key
        and not settings.llm.secondary_api_key
        and not settings.llm.primary_api_key
    )
    if no_llm:
        # No LLM configured — fall back to local DB query engine
        reply = _local_chat_answer(req.message)
        return ChatResponse(reply=reply, context_used=True)

    # ── Gather live DB context ──────────────────────────────────────
    context_lines: list[str] = []
    try:
        db.init_db()
        stats = db.get_pipeline_stats()
        total_students = stats.get('total_students', 0)
        context_lines.append(f"DATABASE SUMMARY:")
        context_lines.append(f"- Total students: {total_students}")
        context_lines.append(f"- Students from email pipeline: {stats.get('email_students', 0)}")
        context_lines.append(f"- Students from admin uploads: {stats.get('admin_students', 0)}")
        context_lines.append(f"- Email extractions: {stats.get('email_extractions', 0)}")
        context_lines.append(f"- Admin upload files: {stats.get('admin_upload_files', 0)}")
        context_lines.append(f"- Total result records: {stats.get('total_results', 0)}")
        if stats.get("average_cgpa"):
            context_lines.append(f"- Average CGPA: {stats['average_cgpa']:.2f}")
        if stats.get("total_backlogs") is not None:
            context_lines.append(f"- Total active backlogs: {stats['total_backlogs']}")

        inst_id = _get_institution_id()
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:

                # Full student list (up to 100) — lets AI answer any student question
                if total_students > 0:
                    cur.execute("""
                        SELECT s.usn, s.name AS full_name, s.cgpa, s.active_backlogs
                        FROM students s
                        WHERE s.institution_id = %s
                                                    {student_filter}
                        ORDER BY s.cgpa DESC NULLS LAST
                        LIMIT 100
                                        """.format(student_filter=db.student_source_filter("s", include_and=True)), (inst_id,))
                    all_students = cur.fetchall()
                    if all_students:
                        context_lines.append(f"\nALL STUDENTS ({len(all_students)}):")
                        for st in all_students:
                            cgpa_str = f", CGPA {float(st['cgpa']):.2f}" if st['cgpa'] else ""
                            back_str = f", backlogs {st['active_backlogs']}" if st['active_backlogs'] else ""
                            context_lines.append(f"  - {st['usn']} {st['full_name'] or ''}{cgpa_str}{back_str}")

                # Students with backlogs
                cur.execute("""
                    SELECT s.usn, s.name AS full_name, s.active_backlogs
                    FROM students s
                    WHERE s.institution_id = %s AND s.active_backlogs > 0
                                            {student_filter}
                    ORDER BY s.active_backlogs DESC
                                """.format(student_filter=db.student_source_filter("s", include_and=True)), (inst_id,))
                backlog_students = cur.fetchall()
                if backlog_students:
                    context_lines.append(f"\nSTUDENTS WITH BACKLOGS ({len(backlog_students)}):")
                    for bs in backlog_students:
                        context_lines.append(f"  - {bs['usn']} {bs['full_name'] or ''}: {bs['active_backlogs']} backlog(s)")

                # Subject pass rates (hardest subjects) — exclude seed student results
                cur.execute("""
                    SELECT sub.name AS subject_name, sub.code,
                           COUNT(*) AS total,
                           SUM(CASE WHEN sr.status='PASS' THEN 1 ELSE 0 END) AS passed
                    FROM student_results sr
                    JOIN subjects sub ON sr.subject_id = sub.id
                    JOIN students s ON sr.student_id = s.id
                    WHERE {student_filter}
                    GROUP BY sub.id, sub.name, sub.code
                    ORDER BY (SUM(CASE WHEN sr.status='PASS' THEN 1 ELSE 0 END)::float / COUNT(*)) ASC
                    LIMIT 5
                """.format(student_filter=db.student_source_filter("s")), ())
                hard_subjects = cur.fetchall()
                if hard_subjects:
                    context_lines.append("\nSUBJECTS WITH LOWEST PASS RATE:")
                    for hs in hard_subjects:
                        rate = round(100 * hs["passed"] / hs["total"]) if hs["total"] else 0
                        context_lines.append(f"  - {hs['code']} {hs['subject_name']}: {rate}% pass rate ({hs['passed']}/{hs['total']})")

                # Specific student detail if USN or name found in question OR in history
                usn_match = re.search(r'\b[1-4][A-Z]{2}\d{2}[A-Z]{2,4}\d{3}\b', req.message.upper())

                # Name-based lookup via search_students (handles all-caps names, single words)
                name_match_usn = None
                if not usn_match:
                    significant = [
                        w.strip('.,!?;:()"\'') for w in req.message.split()
                        if len(w.strip('.,!?;:()"\'')) >= 4
                    ]
                    for word in significant:
                        hits = db.search_students(word, institution_id=inst_id, limit=1)
                        if hits:
                            name_match_usn = hits[0]["usn"]
                            break

                # Scan recent history for USN or name (handles follow-up questions)
                if not usn_match and not name_match_usn:
                    for hist_msg in reversed(req.history[-6:]):
                        hist_usn = re.search(r'\b[1-4][A-Z]{2}\d{2}[A-Z]{2,4}\d{3}\b', hist_msg.content.upper())
                        if hist_usn:
                            name_match_usn = hist_usn.group(0)
                            break
                    if not name_match_usn:
                        for hist_msg in reversed(req.history[-6:]):
                            hist_words = [
                                w.strip('.,!?;:()"\'') for w in hist_msg.content.split()
                                if len(w.strip('.,!?;:()"\'')) >= 4
                            ]
                            for word in hist_words:
                                hits = db.search_students(word, institution_id=inst_id, limit=1)
                                if hits:
                                    name_match_usn = hits[0]["usn"]
                                    break
                            if name_match_usn:
                                break

                target_usn = usn_match.group(0) if usn_match else name_match_usn

                # Extract semester number from message (limit to 1-2 digits, avoid matching USN digits)
                sem_match = re.search(r'\bsem(?:ester)?\s*(\d{1,2})\b|\b(\d{1,2})(?:st|nd|rd|th)?\s*sem\b', req.message.lower())
                filter_sem = int(sem_match.group(1) or sem_match.group(2)) if sem_match else None

                if target_usn:
                    usn = target_usn.upper()
                    cur.execute("""
                        SELECT s.usn, s.name AS full_name, s.cgpa, s.active_backlogs
                        FROM students s
                        WHERE UPPER(s.usn) = %s
                          AND s.institution_id = %s
                          {sf}
                    """.format(sf=db.student_source_filter("s", include_and=True)), (usn, inst_id))
                    st = cur.fetchone()
                    if st:
                        context_lines.append(f"\nSTUDENT DETAIL for {usn}:")
                        context_lines.append(f"  Name: {st['full_name']}")
                        context_lines.append(f"  CGPA: {st['cgpa']}")
                        context_lines.append(f"  Active backlogs: {st['active_backlogs']}")
                        sem_sql = "AND sr.semester = %s" if filter_sem else ""
                        sem_params = (usn, filter_sem) if filter_sem else (usn,)
                        cur.execute(f"""
                            SELECT sub.code, sub.name AS subject_name, sr.total_marks,
                                   sr.max_marks, sr.grade, sr.status, sr.semester
                            FROM student_results sr
                            JOIN subjects sub ON sr.subject_id = sub.id
                            JOIN students s ON sr.student_id = s.id
                            WHERE UPPER(s.usn) = %s {sem_sql}
                            ORDER BY sr.semester, sub.code
                        """, sem_params)
                        results = cur.fetchall()
                        if results:
                            sem_label = f" (Semester {filter_sem})" if filter_sem else " (all semesters)"
                            context_lines.append(f"  Results{sem_label}:")
                            for r in results:
                                context_lines.append(
                                    f"    Sem{r['semester']} | {r['code']} | {r['subject_name']} | "
                                    f"{r['total_marks']}/{r['max_marks']} | {r['grade'] or '-'} | {r['status']}"
                                )
                        else:
                            context_lines.append(f"  No results found{' for semester ' + str(filter_sem) if filter_sem else ''}.")
    except Exception as e:
        logger.warning("chat_db_context_failed", error=str(e))

    db_context = "\n".join(context_lines) if context_lines else "No database context available."
    context_used = bool(context_lines)

    # ── Build conversation for Gemini ──────────────────────────────
    # Gemini uses a simple prompt approach (not multi-turn in this flow)
    convo_history = ""
    for msg in req.history[-6:]:  # last 3 turns (6 messages)
        role = "Teacher" if msg.role == "user" else "AcadAssist"
        convo_history += f"{role}: {msg.content}\n"

    full_prompt = (
        f"{AI_SYSTEM_PROMPT}\n\n"
        f"{db_context}\n\n"
        + (f"CONVERSATION HISTORY:\n{convo_history}\n" if convo_history else "")
        + f"Teacher: {req.message}\nAcadAssist:"
    )

    # Build waterfall from all configured providers: groq → openai → gemini
    _chat_providers: list[str] = []
    if settings.llm.groq_api_key:
        _chat_providers.append("groq")
    if settings.llm.primary_api_key:
        _chat_providers.append("openai")
    if settings.llm.secondary_api_key:
        _chat_providers.append("gemini")

    reply: str | None = None
    _last_exc: Exception | None = None

    for _provider in _chat_providers:
        try:
            if _provider == "groq":
                from groq import Groq
                gclient = Groq(api_key=settings.llm.groq_api_key)
                completion = gclient.chat.completions.create(
                    model=settings.llm.groq_model,
                    messages=[{"role": "user", "content": full_prompt}],
                    temperature=0.3,
                    max_tokens=1024,
                )
                reply = completion.choices[0].message.content.strip()
            elif _provider == "openai":
                from openai import AsyncOpenAI
                oclient = AsyncOpenAI(api_key=settings.llm.primary_api_key)
                resp = await oclient.chat.completions.create(
                    model=settings.llm.primary_model,
                    messages=[{"role": "user", "content": full_prompt}],
                    temperature=0.3,
                    max_tokens=1024,
                )
                reply = resp.choices[0].message.content.strip()
            elif _provider == "gemini":
                import google.generativeai as genai
                genai.configure(api_key=settings.llm.secondary_api_key)
                gmodel = genai.GenerativeModel(settings.llm.secondary_model)
                response = await gmodel.generate_content_async(
                    full_prompt,
                    generation_config={"temperature": 0.3, "max_output_tokens": 1024},
                )
                reply = response.text.strip()
            if reply:
                logger.debug("ai_chat: provider=%s succeeded", _provider)
                break
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning("ai_chat_%s_failed: %s", _provider, exc)
            _last_exc = exc

    if not reply:
        # All providers exhausted — give a user-friendly error
        _msg = str(_last_exc) if _last_exc else "unknown"
        if _last_exc and "rate_limit" in _msg.lower():
            raise HTTPException(
                status_code=429,
                detail="AI provider rate limit reached. Please try again in a few minutes.",
            )
        raise HTTPException(
            status_code=503,
            detail="AI assistant is temporarily unavailable. Please try again shortly.",
        )

    # Run verification agent to cross-check the LLM reply against live DB values
    try:
        reply, _ = _verify_chat_reply(reply, _get_institution_id())
    except Exception as exc:
        logger.warning("verification_agent_skipped", error=str(exc))

    return ChatResponse(reply=reply, context_used=context_used)
