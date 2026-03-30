"""
Admin Endpoints.

System statistics, pipeline management, and document upload for administrators.
"""

from __future__ import annotations

import re
import io
from typing import Any

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel

from src.common.observability import get_logger
from src.common import database as db

logger = get_logger(__name__)
router = APIRouter()


# ── Endpoints ───────────────────────────────────────────────────────


@router.get("/status")
async def pipeline_status() -> dict:
    """Get current pipeline status and database statistics."""
    try:
        db.init_db()
        stats = db.get_pipeline_stats()
        return {"status": "ok", "stats": stats}
    except Exception as e:
        logger.error("admin_status_failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats")
async def system_statistics() -> dict:
    """Get system-wide academic statistics."""
    try:
        db.init_db()
        stats = db.get_pipeline_stats()
        return {
            "total_students": stats["total_students"],
            "students_from_email": stats["email_students"],
            "students_from_admin": stats["admin_students"],
            "total_results": stats["total_results"],
            "email_extractions": stats["email_extractions"],
            "admin_upload_files": stats["admin_upload_files"],
            "avg_cgpa": stats["average_cgpa"],
            "students_with_backlogs": stats["total_backlogs"],
        }
    except Exception as e:
        logger.error("admin_stats_failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ── Document Upload ──────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_document(file: UploadFile = File(...)) -> dict:
    """
    Upload an Excel, CSV, PDF, DOCX, or TXT result document.
    Parses student records and stores them directly to the database.
    Returns a summary of what was extracted.
    """
    filename = (file.filename or "").lower()
    content  = await file.read()

    db.init_db()
    inst_id = db.get_default_institution_id()

    try:
        if filename.endswith((".xlsx", ".xls")):
            records = _parse_excel(content, filename)
        elif filename.endswith(".csv"):
            records = _parse_csv(content)
        elif filename.endswith(".pdf"):
            records = _parse_pdf_text(content)
        elif filename.endswith(".docx"):
            records = _parse_docx_text(content)
        elif filename.endswith((".txt", ".text")):
            records = _parse_raw_text(content.decode("utf-8", errors="ignore"))
        elif filename.endswith(".zip"):
            records = _parse_zip(content)
        else:
            raise HTTPException(
                status_code=415,
                detail="Unsupported file type. Upload Excel (.xlsx), CSV, PDF, DOCX, TXT, or ZIP."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("document_parse_failed", filename=filename, error=str(e))
        raise HTTPException(status_code=422, detail=f"Failed to parse file: {e}")

    if not records:
        return {
            "status": "no_data",
            "message": "No student records found in this file. Ensure it contains USN and result data.",
            "filename": file.filename,
        }

    # Store records
    students_upserted = 0
    results_stored    = 0
    errors: list[str] = []

    for rec in records:
        try:
            usn  = str(rec.get("usn", "")).strip().upper()
            name = str(rec.get("name", usn)).strip()
            if not usn:
                continue

            student_id = db.upsert_student(usn, name, institution_id=inst_id, source="upload")
            students_upserted += 1

            # Semester aggregate (SGPA-only records)
            if rec.get("sgpa") and not rec.get("subject_code"):
                sem = int(rec.get("semester", 1))
                sgpa = float(rec["sgpa"])
                cgpa = float(rec.get("cgpa", 0)) or None
                db.store_semester_aggregate(student_id=student_id, semester=sem, sgpa=sgpa, backlogs=0)
                if cgpa:
                    with db.get_connection() as conn:
                        with conn.cursor() as cur:
                            cur.execute("UPDATE students SET cgpa=%s WHERE id=%s", (cgpa, student_id))

            # Full subject result
            elif rec.get("subject_code"):
                sem     = int(rec.get("semester", 1))
                code    = str(rec.get("subject_code", "UNKNOWN")).strip()
                subname = str(rec.get("subject_name", code)).strip()
                marks   = float(rec.get("marks", 0) or 0)
                max_m   = float(rec.get("max_marks", 100) or 100)
                grade   = str(rec.get("grade", "")).strip() or None
                status  = str(rec.get("status", "PASS" if marks >= 40 else "FAIL")).upper()
                gp      = float(rec.get("grade_points", 0) or 0) or None

                subject_id = db.get_or_create_subject(inst_id, code, subname, sem)
                db.upsert_result(
                    student_id=student_id,
                    subject_id=subject_id,
                    semester=sem,
                    marks_obtained=marks,
                    max_marks=max_m,
                    grade=grade,
                    grade_points=gp,
                    status=status,
                )
                results_stored += 1

            db.compute_and_store_cgpa(student_id)

        except Exception as e:
            errors.append(f"USN {rec.get('usn','?')}: {e}")

    db.save_admin_upload(
        institution_id=inst_id,
        filename=file.filename or "upload",
        content_type=file.content_type,
        file_size=len(content),
        records_parsed=len(records),
        students_upserted=students_upserted,
        results_stored=results_stored,
    )

    return {
        "status": "ok",
        "filename": file.filename,
        "students_upserted": students_upserted,
        "results_stored": results_stored,
        "records_parsed": len(records),
        "errors": errors[:10],  # cap error list
    }

# ── Parsing helpers ───────────────────────────────────────────────────────────

_USN_RE  = re.compile(r"\b([1-4][A-Z]{2}\d{2}[A-Z]{2,4}\d{3})\b", re.I)
_SGPA_RE = re.compile(r"(?:s\.?g\.?p\.?a|semester\s+gpa)[^\d]*([\d]{1,2}\.[\d]{1,2})", re.I)
_CGPA_RE = re.compile(r"(?:c\.?g\.?p\.?a|cumulative\s+gpa|overall\s+gpa)[^\d]*([\d]{1,2}\.[\d]{1,2})", re.I)


def _parse_raw_text(text: str) -> list[dict]:
    """Extract student records from free-form text using regex."""
    records: list[dict] = []
    # Split into per-student blocks at each USN
    blocks = re.split(r"(?=\b[1-4][A-Z]{2}\d{2}[A-Z]{2,4}\d{3}\b)", text, flags=re.I)
    for block in blocks:
        usn_m = _USN_RE.search(block)
        if not usn_m:
            continue
        usn  = usn_m.group(1).upper()
        sgpa = _SGPA_RE.search(block)
        cgpa = _CGPA_RE.search(block)
        sem_m = re.search(r"semester[:\s]+(\d)", block, re.I)
        sem = int(sem_m.group(1)) if sem_m else 1
        # Name heuristic: first ≥2 capitalized words after USN
        post = block[usn_m.end():].strip()
        name_m = re.match(r"([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3})", post)
        name = name_m.group(1).strip() if name_m else usn
        rec: dict[str, Any] = {"usn": usn, "name": name, "semester": sem}
        if sgpa:
            rec["sgpa"] = float(sgpa.group(1))
        if cgpa:
            rec["cgpa"] = float(cgpa.group(1))
        # Subject results: CODE  Name  Marks  Grade  PASS/FAIL
        subj_re = re.compile(
            r"\b([A-Z]{2,4}\d{2,4})\b[^\n]*?(\d{2,3})\s*/?\s*(\d{2,3})?\s+([A-F][+-]?)\s+(PASS|FAIL)",
            re.I
        )
        subs = subj_re.findall(block)
        if subs:
            for code, marks, max_m, grade, status in subs:
                records.append({
                    "usn": usn, "name": name, "semester": sem,
                    "subject_code": code.upper(), "subject_name": code.upper(),
                    "marks": float(marks), "max_marks": float(max_m or 100),
                    "grade": grade.upper(), "status": status.upper(),
                })
        else:
            records.append(rec)
    return records


def _parse_pdf_text(content: bytes) -> list[dict]:
    import pdfplumber
    import pandas as pd

    text_parts: list[str] = []
    table_records: list[dict] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # ── 1. Try table extraction first (grade cards / mark sheets) ──
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue
                try:
                    raw_headers = table[0]
                    headers = [
                        str(c).strip().replace("\n", " ") if c else f"col_{i}"
                        for i, c in enumerate(raw_headers)
                    ]
                    rows = [
                        [str(cell).strip() if cell is not None else "" for cell in row]
                        for row in table[1:]
                    ]
                    df = pd.DataFrame(rows, columns=headers)
                    recs = _parse_dataframe(df)
                    if recs:
                        table_records.extend(recs)
                except Exception:
                    pass

            # ── 2. Also collect plain text for USN/SGPA/CGPA header info ──
            t = page.extract_text()
            if t:
                text_parts.append(t)

    # If table extraction found subject-level records, use them
    if table_records:
        # Enrich with header-level USN/SGPA/CGPA from text if missing
        text_all = "\n".join(text_parts)
        usn_m  = _USN_RE.search(text_all)
        sgpa_m = _SGPA_RE.search(text_all)
        cgpa_m = _CGPA_RE.search(text_all)
        for rec in table_records:
            if not rec.get("usn") and usn_m:
                rec["usn"] = usn_m.group(1).upper()
            if not rec.get("sgpa") and sgpa_m:
                rec["sgpa"] = float(sgpa_m.group(1))
            if not rec.get("cgpa") and cgpa_m:
                rec["cgpa"] = float(cgpa_m.group(1))
        # Filter out any records still missing a valid USN
        valid = [r for r in table_records if r.get("usn") and _USN_RE.match(str(r["usn"]))]
        if valid:
            return valid

    # ── Fallback: regex over plain-text ──────────────────────────────────────
    return _parse_raw_text("\n".join(text_parts))


def _parse_docx_text(content: bytes) -> list[dict]:
    from docx import Document
    doc = Document(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + "\t".join(cell.text for cell in row.cells)
    return _parse_raw_text(text)


def _parse_csv(content: bytes) -> list[dict]:
    import pandas as pd
    try:
        df = pd.read_csv(io.BytesIO(content), dtype=str)
    except Exception:
        df = pd.read_csv(io.BytesIO(content), dtype=str, encoding="latin-1")
    return _parse_dataframe(df)


def _parse_excel(content: bytes, filename: str) -> list[dict]:
    import pandas as pd
    engine = "xlrd" if filename.endswith(".xls") else "openpyxl"
    try:
        df = pd.read_excel(io.BytesIO(content), dtype=str, engine=engine)
    except Exception:
        # Try reading all sheets and pick the first non-empty one
        sheets = pd.read_excel(io.BytesIO(content), dtype=str, engine=engine, sheet_name=None)
        df = next((v for v in sheets.values() if not v.empty), None)
        if df is None:
            return []
    records = _parse_dataframe(df)
    if records:
        return records
    # Fallback: VTU-style wide-format grade report (no standard column headers)
    return _parse_vtu_grade_report(content, engine)


def _parse_vtu_grade_report(content: bytes, engine: str) -> list[dict]:
    """
    Parse VTU-style wide-format grade reports where:
      - Row ~2: Sl No. | USN | Name | SUBJ_CODE | nan | nan | SUBJ_CODE | nan | ...
      - Row ~5: GR | nan | GP | GR | GP | GR | GP | ...  (grade / grade-points per subject)
      - Row 6+: student data rows
    """
    import pandas as pd

    raw = pd.read_excel(io.BytesIO(content), dtype=str, header=None, engine=engine)
    n_cols = len(raw.columns)
    n_rows = len(raw)

    # Locate the header row that contains a cell exactly equal to "USN"
    hdr_row = None
    for i in range(min(20, n_rows)):
        for c in range(n_cols):
            if str(raw.iloc[i, c]).strip().upper() == "USN":
                hdr_row = i
                break
        if hdr_row is not None:
            break
    if hdr_row is None:
        return []

    row_hdr = [str(raw.iloc[hdr_row, c]).strip() for c in range(n_cols)]

    usn_col  = next((c for c, v in enumerate(row_hdr) if v.upper() == "USN"), None)
    name_col = next((c for c, v in enumerate(row_hdr) if v.upper() in ("NAME", "STUDENT NAME", "FULL NAME")), None)
    if usn_col is None:
        return []

    # Subject names live one row below the header row
    if hdr_row + 1 < n_rows:
        row_snames = [str(raw.iloc[hdr_row + 1, c]).strip() for c in range(n_cols)]
    else:
        row_snames = ["nan"] * n_cols

    # Locate the GR/GP sub-header row (first row after hdr_row containing both "GR" and "GP")
    gr_gp_row = None
    for i in range(hdr_row + 1, min(hdr_row + 8, n_rows)):
        vals = [str(raw.iloc[i, c]).strip().upper() for c in range(n_cols)]
        if "GR" in vals and "GP" in vals:
            gr_gp_row = i
            break

    # Non-subject columns to skip
    meta_cols: set[int] = {usn_col}
    if name_col is not None:
        meta_cols.add(name_col)
    for c, v in enumerate(row_hdr):
        if v.upper() in ("SL NO.", "SL NO", "SLNO", "S.NO.", "S.NO", "SNO", ""):
            meta_cols.add(c)

    # Build subject list: (code, subject_name, grade_col, gp_col)
    subjects: list[tuple[str, str, int, int | None]] = []
    for c in range(n_cols):
        if c in meta_cols:
            continue
        code = row_hdr[c]
        if code.lower() in ("nan", ""):
            continue  # merged/filler cell

        # Determine subject name from the names row
        sname = row_snames[c] if row_snames[c].lower() not in ("nan", "") else code
        sname = sname.replace("\n", " ")

        # Find the GR column (= c itself if gr_gp_row says "GR" there, else scan forward)
        gr_col: int | None = None
        gp_col: int | None = None
        if gr_gp_row is not None:
            for ci in range(c, min(c + 4, n_cols)):
                v = str(raw.iloc[gr_gp_row, ci]).strip().upper()
                if v == "GR" and gr_col is None:
                    gr_col = ci
                elif v == "GP" and gr_col is not None:
                    gp_col = ci
                    break
        else:
            gr_col, gp_col = c, c + 1 if c + 1 < n_cols else None

        if gr_col is not None:
            subjects.append((code, sname, gr_col, gp_col))

    if not subjects:
        return []

    # Infer semester from metadata rows before hdr_row
    sem = 1
    roman_map = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8}
    for i in range(hdr_row):
        text = " ".join(
            str(raw.iloc[i, c]) for c in range(n_cols)
            if str(raw.iloc[i, c]).strip().lower() not in ("nan", "")
        )
        m = re.search(r"[Ss]emester\s*:?\s*([IVXivx]+|\d+)", text)
        if m:
            t = m.group(1).strip().upper()
            sem = roman_map.get(t, int(t) if t.isdigit() else 1)
            break

    data_start = (gr_gp_row + 1) if gr_gp_row is not None else hdr_row + 4
    records: list[dict] = []

    for row_i in range(data_start, n_rows):
        usn_val = str(raw.iloc[row_i, usn_col]).strip().upper()
        m = _USN_RE.search(usn_val)
        if not m:
            continue
        usn = m.group(1)

        name = str(raw.iloc[row_i, name_col]).strip() if name_col is not None else usn
        if name.lower() in ("nan", ""):
            name = usn

        for code, sname, gr_col, gp_col in subjects:
            grade = str(raw.iloc[row_i, gr_col]).strip() if gr_col < n_cols else ""
            if grade.lower() in ("nan", "", "ne", "ab"):
                continue  # not registered / absent / not eligible

            gp_raw = str(raw.iloc[row_i, gp_col]).strip() if gp_col is not None and gp_col < n_cols else ""
            gp: float | None = None
            try:
                if gp_raw.lower() not in ("nan", ""):
                    gp = float(gp_raw.replace(",", "."))
            except (ValueError, TypeError):
                pass

            status = "FAIL" if grade.upper() == "F" or (gp is not None and gp == 0) else "PASS"
            records.append({
                "usn": usn,
                "name": name,
                "semester": sem,
                "subject_code": code,
                "subject_name": sname,
                "grade": grade,
                "grade_points": gp,
                "status": status,
            })

    return records


def _parse_dataframe(df) -> list[dict]:
    """
    Parse a DataFrame where each row may be a student+semester aggregate
    or a student+subject result.

    Column name aliases (case-insensitive) supported:
      USN column    : usn, roll, roll_no, roll_number, university_seat_number, id
      Name column   : name, student_name, full_name
      Semester col  : semester, sem, semno
      SGPA col      : sgpa, semester_gpa, sem_gpa
      CGPA col      : cgpa
      Subject code  : subject_code, code, sub_code
      Subject name  : subject, subject_name, sub_name, paper
      Marks col     : marks, marks_obtained, total_marks, score
      Max marks     : max_marks, max, total, out_of
      Grade col     : grade, letter_grade
      Status col    : status, result, pass_fail
      Grade points  : grade_points, gp, points
    """
    import pandas as pd

    # Normalize column names
    df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]

    def _pick(cols: list[str]):
        for c in cols:
            if c in df.columns:
                return c
        return None

    usn_col    = _pick(["usn","roll","roll_no","roll_number","university_seat_number","id","seat_no"])
    name_col   = _pick(["name","student_name","full_name","student"])
    sem_col    = _pick(["semester","sem","semno","sem_no"])
    sgpa_col   = _pick(["sgpa","semester_gpa","sem_gpa","gpa"])
    cgpa_col   = _pick(["cgpa","cumulative_gpa"])
    code_col   = _pick(["subject_code","code","sub_code","course_code"])
    sname_col  = _pick(["subject","subject_name","sub_name","paper","course_name","course"])
    marks_col  = _pick(["marks","marks_obtained","total_marks","score","obtained"])
    max_col    = _pick(["max_marks","max","total","out_of","full_marks"])
    grade_col  = _pick(["grade","letter_grade","grade_letter"])
    status_col = _pick(["status","result","pass_fail","outcome"])
    gp_col     = _pick(["grade_points","gp","points","grade_point"])

    if not usn_col:
        return []

    records: list[dict] = []
    for _, row in df.iterrows():
        usn = str(row.get(usn_col, "")).strip().upper()
        # Validate USN format
        if not _USN_RE.match(usn):
            # Try to find USN pattern anywhere in the column value
            m = _USN_RE.search(usn)
            if m:
                usn = m.group(1)
            else:
                continue

        name = str(row.get(name_col, usn)).strip() if name_col else usn
        sem  = int(float(row.get(sem_col, 1) or 1)) if sem_col else 1
        rec: dict[str, Any] = {"usn": usn, "name": name, "semester": sem}

        if sgpa_col and pd.notna(row.get(sgpa_col)):
            try:
                rec["sgpa"] = float(str(row[sgpa_col]).replace(",", "."))
            except (ValueError, TypeError):
                pass

        if cgpa_col and pd.notna(row.get(cgpa_col)):
            try:
                rec["cgpa"] = float(str(row[cgpa_col]).replace(",", "."))
            except (ValueError, TypeError):
                pass

        if code_col and pd.notna(row.get(code_col)):
            rec["subject_code"] = str(row[code_col]).strip().upper()
            if sname_col and pd.notna(row.get(sname_col)):
                rec["subject_name"] = str(row[sname_col]).strip()
            else:
                rec["subject_name"] = rec["subject_code"]
            if marks_col and pd.notna(row.get(marks_col)):
                try:
                    rec["marks"] = float(str(row[marks_col]).replace(",","."))
                except (ValueError, TypeError):
                    pass
            if max_col and pd.notna(row.get(max_col)):
                try:
                    rec["max_marks"] = float(str(row[max_col]).replace(",","."))
                except (ValueError, TypeError):
                    pass
            if grade_col and pd.notna(row.get(grade_col)):
                rec["grade"] = str(row[grade_col]).strip()
            if status_col and pd.notna(row.get(status_col)):
                raw_s = str(row[status_col]).strip().upper()
                rec["status"] = "PASS" if raw_s in ("PASS","P","1","TRUE","YES","OK") else "FAIL"
            if gp_col and pd.notna(row.get(gp_col)):
                try:
                    rec["grade_points"] = float(str(row[gp_col]).replace(",","."))
                except (ValueError, TypeError):
                    pass

        records.append(rec)
    return records


def _parse_zip(content: bytes) -> list[dict]:
    """
    Extract files from a ZIP archive and parse each supported file.
    Returns combined records from all extracted files.
    """
    import zipfile

    records: list[dict] = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            for name in zf.namelist():
                # Skip __MACOSX metadata, hidden files, and directories
                base = name.rsplit("/", 1)[-1]
                if not base or base.startswith(".") or base.startswith("__"):
                    continue
                lower = base.lower()
                try:
                    file_bytes = zf.read(name)
                    if lower.endswith((".xlsx", ".xls")):
                        recs = _parse_excel(file_bytes, lower)
                    elif lower.endswith(".csv"):
                        recs = _parse_csv(file_bytes)
                    elif lower.endswith(".pdf"):
                        recs = _parse_pdf_text(file_bytes)
                    elif lower.endswith(".docx"):
                        recs = _parse_docx_text(file_bytes)
                    elif lower.endswith((".txt", ".text")):
                        recs = _parse_raw_text(file_bytes.decode("utf-8", errors="ignore"))
                    else:
                        continue  # skip unsupported files inside zip
                    records.extend(recs)
                    logger.info("zip_file_parsed", inner_file=name, records=len(recs))
                except Exception as exc:
                    logger.warning("zip_inner_file_failed", inner_file=name, error=str(exc))
    except zipfile.BadZipFile:
        raise HTTPException(status_code=422, detail="Invalid or corrupted ZIP file.")
    return records


# ── Review Queue ─────────────────────────────────────────────────────────────


@router.get("/review-queue")
async def list_review_queue(
    status: str = "pending",
    limit: int = 50,
) -> dict:
    """List review queue items filtered by status (pending/approved/rejected/corrected)."""
    from src.phase3_extraction_engine.review_queue import get_review_queue
    if status not in ("pending", "approved", "rejected", "corrected"):
        raise HTTPException(status_code=400, detail="status must be pending|approved|rejected|corrected")
    items = get_review_queue(status=status, limit=min(limit, 200))
    return {"status": "ok", "count": len(items), "items": items}


class ReviewApproveRequest(BaseModel):
    corrected_data: list[dict] | None = None
    notes: str = ""
    save_to_db: bool = True


class ReviewRejectRequest(BaseModel):
    notes: str = ""


@router.post("/review-queue/{item_id}/approve")
async def approve_review_item(item_id: str, body: ReviewApproveRequest) -> dict:
    """
    Approve a review queue item.

    If save_to_db=True (default) the approved/corrected records are persisted to
    the students and student_results tables. If corrected_data is provided those
    values are used; otherwise the original extracted_data is saved.
    """
    from src.phase3_extraction_engine.review_queue import (
        approve_review_item as _approve,
        get_review_queue,
    )

    # Fetch the item first so we can save records after approval
    items = None
    if body.save_to_db:
        all_items = get_review_queue(status="pending", limit=500)
        items = next((i for i in all_items if str(i.get("id")) == item_id), None)

    ok = _approve(item_id, corrected_data=body.corrected_data, notes=body.notes)
    if not ok:
        raise HTTPException(status_code=404, detail="Review queue item not found or update failed.")

    saved = 0
    if body.save_to_db and items:
        try:
            import json as _json
            records_to_save = body.corrected_data
            if not records_to_save:
                raw = items.get("extracted_data")
                if isinstance(raw, str):
                    records_to_save = _json.loads(raw)
                elif isinstance(raw, list):
                    records_to_save = raw

            if records_to_save:
                inst_id = db.get_default_institution_id()
                for rec in records_to_save:
                    usn = rec.get("usn")
                    if not usn:
                        continue
                    student_id = db.upsert_student(
                        usn, rec.get("name", usn), institution_id=inst_id, source="upload"
                    )
                    for subj in rec.get("subjects", []):
                        subj_code = subj.get("subject_code", "UNKNOWN")
                        subj_name = subj.get("subject_name", subj_code)
                        subject_id = db.get_or_create_subject(
                            inst_id, subj_code, subj_name, rec.get("semester")
                        )
                        db.upsert_result(
                            student_id=student_id,
                            subject_id=subject_id,
                            semester=rec.get("semester", 1),
                            marks_obtained=subj.get("total_marks"),
                            max_marks=subj.get("max_marks", 100),
                            grade=subj.get("grade"),
                            grade_points=subj.get("grade_points"),
                            status=subj.get("status"),
                        )
                        saved += 1
                    student = db.get_student(usn, inst_id)
                    if student:
                        db.compute_and_store_cgpa(student["id"])
        except Exception as exc:
            logger.warning("review_queue.approve: db save failed: %s", exc)

    return {"status": "ok", "item_id": item_id, "records_saved": saved}


@router.post("/review-queue/{item_id}/reject")
async def reject_review_item(item_id: str, body: ReviewRejectRequest) -> dict:
    """Reject a review queue item."""
    from src.phase3_extraction_engine.review_queue import reject_review_item as _reject

    ok = _reject(item_id, notes=body.notes)
    if not ok:
        raise HTTPException(status_code=404, detail="Review queue item not found or update failed.")
    return {"status": "ok", "item_id": item_id}


@router.post("/fix-grades")
async def fix_grades() -> dict:
    """
    One-time data migration: fixes corrupted grade/grade_points values from
    earlier pipeline runs and recomputes SGPA/CGPA for all students.

    Run once after upgrading to the new strict-validation pipeline.
    Safe to call multiple times (idempotent SQL CASE expressions).
    """
    try:
        stats = db.fix_corrupted_grade_data()
        return {"status": "ok", "migration": stats}
    except Exception as exc:
        logger.error("fix_grades_failed", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc))


# ── Institution config ─────────────────────────────────────────────────────────

class InstitutionConfigUpdate(BaseModel):
    config: dict


@router.get("/institution/config")
async def get_institution_config(institution_id: str | None = None) -> dict:
    """
    Return the JSONB config for the institution.
    Supported keys:
      - pass_threshold (int): minimum marks for PASS (default: 35)
      - (extensible for future institution-level settings)
    """
    try:
        cfg = db.get_institution_config(institution_id)
        return {"status": "ok", "config": cfg}
    except Exception as exc:
        logger.error("get_institution_config_failed", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc))


@router.patch("/institution/config")
async def patch_institution_config(
    body: InstitutionConfigUpdate,
    institution_id: str | None = None,
) -> dict:
    """
    Merge-update the institution JSONB config.
    Only keys provided in the request body are updated; other keys are preserved.

    Example body: {"config": {"pass_threshold": 40}}
    """
    try:
        db.set_institution_config(body.config, institution_id)
        cfg = db.get_institution_config(institution_id)
        return {"status": "ok", "config": cfg}
    except Exception as exc:
        logger.error("patch_institution_config_failed", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc))
