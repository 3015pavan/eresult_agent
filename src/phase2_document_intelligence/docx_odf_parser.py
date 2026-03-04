"""
DOCX / ODT / ODF / RTF Parser — Phase 2.

Handles:
  .docx            — python-docx (primary)
  .odt / .odf      — odfpy
  .rtf             — striprtf / regex fallback
  .txt / .md       — plain read

Returns ParsedDocument with extracted text and any detected tables.
"""

from __future__ import annotations

import logging
import os
import re
from typing import Optional

from .router import ParsedDocument

logger = logging.getLogger(__name__)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(path: str) -> tuple[str, list[list[list[str]]]]:
    """Extract text + tables from a .docx file using python-docx."""
    import docx as python_docx  # type: ignore

    doc  = python_docx.Document(path)
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    for tbl in doc.tables:
        rows: list[list[str]] = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
            # Also add pipe-text version for downstream regex extractor
            for row in rows:
                text_parts.append(" | ".join(row))

    return "\n".join(text_parts), tables


# ── ODT / ODF ─────────────────────────────────────────────────────────────────

def _parse_odt(path: str) -> tuple[str, list[list[list[str]]]]:
    """Extract text + tables from .odt / .odf files using odfpy."""
    from odf import text as odf_text, teletype  # type: ignore
    from odf.opendocument import load as odf_load  # type: ignore
    from odf.table import Table, TableRow, TableCell  # type: ignore

    doc        = odf_load(path)
    text_parts: list[str] = []
    tables: list[list[list[str]]] = []

    # Extract paragraphs
    for para in doc.getElementsByType(odf_text.P):
        line = teletype.extractText(para).strip()
        if line:
            text_parts.append(line)

    # Extract tables
    for tbl in doc.getElementsByType(Table):
        rows: list[list[str]] = []
        for row in tbl.getElementsByType(TableRow):
            cells: list[str] = []
            for cell in row.getElementsByType(TableCell):
                cells.append(teletype.extractText(cell).strip())
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
            for row in rows:
                text_parts.append(" | ".join(row))

    return "\n".join(text_parts), tables


# ── RTF ───────────────────────────────────────────────────────────────────────

def _parse_rtf(path: str) -> str:
    """Strip RTF markup and return plain text."""
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
        with open(path, encoding="utf-8", errors="replace") as f:
            return rtf_to_text(f.read())
    except ImportError:
        pass

    # Regex fallback for basic RTF
    with open(path, encoding="utf-8", errors="replace") as f:
        rtf = f.read()
    rtf = re.sub(r"\{[^{}]*\}", "", rtf)          # strip groups
    rtf = re.sub(r"\\[a-zA-Z]+\d*\s?", " ", rtf) # strip control words
    rtf = re.sub(r"[{}\\]", "", rtf)
    return rtf.strip()


# ── Plain text ────────────────────────────────────────────────────────────────

def _parse_plain(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


# ── Public API ────────────────────────────────────────────────────────────────

def parse_document_file(path: str, mime_type: str = "") -> ParsedDocument:
    """
    Parse a DOCX / ODT / ODF / RTF / TXT file.

    Args:
        path:      Absolute path to the file.
        mime_type: Optional MIME type hint.

    Returns:
        ParsedDocument with text + tables.
    """
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    text  = ""
    tables: list[list[list[str]]] = []
    strategy = "unknown"
    errors: list[str] = []
    confidence = 0.0

    # ── DOCX ──────────────────────────────────────────────────────────────────
    if ext in ("docx", "doc") or "wordprocessingml" in mime_type or "msword" in mime_type:
        try:
            text, tables = _parse_docx(path)
            strategy   = "python_docx"
            confidence = 0.85
        except ImportError:
            errors.append("python-docx not installed")
        except Exception as exc:
            errors.append(f"docx:{exc}")

    # ── ODT / ODF ─────────────────────────────────────────────────────────────
    elif ext in ("odt", "ods", "odp", "odf") or "opendocument" in mime_type:
        try:
            text, tables = _parse_odt(path)
            strategy   = "odfpy"
            confidence = 0.80
        except ImportError:
            errors.append("odfpy not installed")
        except Exception as exc:
            errors.append(f"odf:{exc}")

    # ── RTF ───────────────────────────────────────────────────────────────────
    elif ext == "rtf" or "rtf" in mime_type:
        try:
            text     = _parse_rtf(path)
            strategy = "rtf_strip"
            confidence = 0.65
        except Exception as exc:
            errors.append(f"rtf:{exc}")

    # ── Markdown / Plain text / CSV-like ─────────────────────────────────────
    elif ext in ("txt", "md", "text", "log") or mime_type.startswith("text/"):
        try:
            text     = _parse_plain(path)
            strategy = "plain_text"
            confidence = 0.75
        except Exception as exc:
            errors.append(f"txt:{exc}")

    else:
        # Unknown type — try reading as plain text
        try:
            text = _parse_plain(path)
            strategy  = "unknown_as_text"
            confidence = 0.40
        except Exception as exc:
            errors.append(f"unknown:{exc}")

    if not confidence and errors:
        confidence = 0.0

    return ParsedDocument(
        source_path=path,
        mime_type=mime_type or f"application/{ext}",
        text=text,
        tables=tables,
        parse_strategy=strategy,
        confidence=confidence,
        errors=errors,
    )
